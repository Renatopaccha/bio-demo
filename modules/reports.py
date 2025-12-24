import streamlit as st
import pandas as pd
import os
import zipfile
import pickle
import json
import datetime
import io
from modules.utils import boton_guardar_grafico, boton_guardar_tabla

# ==========================================
# 1. REPORTES (DOCX/PDF)
# ==========================================
def render_reportes():
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.header("📄 Generador de Reportes")
    
    if not st.session_state.get('reporte_items'):
        st.info("Aún no has añadido elementos al reporte. Ve a las otras pestañas y usa los botones '➕ Añadir'.")
    else:
        st.write(f"Tienes {len(st.session_state['reporte_items'])} elementos listos para exportar.")
        
        # Botón para vaciar reporte (Gestión de Memoria)
        if st.button("🗑️ Vaciar Reporte Completo", type="primary"):
            st.session_state['reporte_items'] = []
            st.success("Reporte vaciado correctamente.")
            st.rerun()

        # Vista previa lista
        with st.expander("Ver lista de elementos"):
            for i, item in enumerate(st.session_state['reporte_items']):
                c1, c2 = st.columns([4, 1])
                c1.write(f"{i+1}. {item['titulo']} ({item['tipo']})")
                if c2.button("❌", key=f"del_rep_{i}"):
                    st.session_state['reporte_items'].pop(i)
                    st.rerun()

        # Configuración Reporte
        titulo_rep = st.text_input("Título del Reporte", "Informe Estadístico - BioStat Easy")
        autor_rep = st.text_input("Autor", "Investigador Principal")
        
        c_exp1, c_exp2 = st.columns(2)
        
        # EXPORTAR WORD
        with c_exp1:
            if st.button("📄 Exportar a Word (.docx)"):
                try:
                    from docx import Document
                    from docx.shared import Inches
                    
                    doc = Document()
                    doc.add_heading(titulo_rep, 0)
                    doc.add_paragraph(f"Autor: {autor_rep}")
                    doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d')}")
                    
                    for item in st.session_state['reporte_items']:
                        doc.add_heading(item['titulo'], level=2)
                        
                        if item['tipo'] == 'img':
                            # Guardar temp para añadir
                            item['data'].seek(0)
                            doc.add_picture(item['data'], width=Inches(6))
                        elif item['tipo'] == 'df':
                            df_rep = item['data']
                            t = doc.add_table(rows=1, cols=len(df_rep.columns))
                            t.style = 'Table Grid'
                            hdr_cells = t.rows[0].cells
                            for k, col_name in enumerate(df_rep.columns):
                                hdr_cells[k].text = str(col_name)
                            
                            for index, row in df_rep.iterrows():
                                row_cells = t.add_row().cells
                                for k, val in enumerate(row):
                                    row_cells[k].text = str(val)
                    
                    # Guardar en buffer
                    f_out = io.BytesIO()
                    doc.save(f_out)
                    f_out.seek(0)
                    
                    st.download_button("⬇️ Descargar Word", f_out, "reporte_biostat.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    st.success("Reporte generado.")
                    
                except ImportError:
                    st.error("Librería 'python-docx' no instalada.")
                except Exception as e:
                    st.error(f"Error generando Word: {e}")

        # EXPORTAR PDF (Básico con FPDF)
        with c_exp2:
            if st.button("📄 Exportar a PDF"):
                try:
                    from fpdf import FPDF
                    import tempfile
                    
                    class PDF(FPDF):
                        def header(self):
                            self.set_font('Arial', 'B', 15)
                            self.cell(0, 10, titulo_rep, 0, 1, 'C')
                            self.ln(10)
                    
                    pdf = PDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.cell(0, 10, f"Autor: {autor_rep}", 0, 1)
                    pdf.cell(0, 10, f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d')}", 0, 1)
                    
                    for item in st.session_state['reporte_items']:
                        pdf.add_page()
                        pdf.set_font("Arial", 'B', 14)
                        pdf.cell(0, 10, item['titulo'], 0, 1)
                        
                        if item['tipo'] == 'img':
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                                item['data'].seek(0)
                                tmp_img.write(item['data'].read())
                                tmp_path = tmp_img.name
                            
                            pdf.image(tmp_path, x=10, w=190)
                            os.unlink(tmp_path)
                            
                        elif item['tipo'] == 'df':
                            pdf.set_font("Arial", size=10)
                            # Tabla simple texto
                            txt_table = item['data'].to_string()
                            pdf.multi_cell(0, 5, txt_table)
                    
                    pdf_out = pdf.output(dest='S').encode('latin-1')
                    st.download_button("⬇️ Descargar PDF", pdf_out, "reporte_biostat.pdf", "application/pdf")
                    
                except ImportError:
                    st.error("Librería 'fpdf' no instalada.")
                except Exception as e:
                    st.error(f"Error generando PDF: {e}")

    st.markdown('</div>', unsafe_allow_html=True)

# ==========================================
# 2. GESTIÓN DE PROYECTOS (ZIP/PARQUET)
# ==========================================
def render_proyectos():
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.header("📂 Gestión de Proyectos (Backup)")
    
    tab1, tab2 = st.tabs(["💾 Guardar Proyecto", "📂 Cargar Proyecto"])
    
    # --- GUARDAR ---
    with tab1:
        st.write("Guarda todo tu trabajo (Datos + Reporte) en un archivo .zip comprimido.")
        if st.button("📦 Generar Backup Completo"):
            try:
                # Crear buffer ZIP en memoria
                zip_buffer = io.BytesIO()
                
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    # 1. Guardar DF Principal (PARQUET por seguridad/velocidad)
                    if st.session_state.df_principal is not None:
                        # Usar Parquet en lugar de Pickle
                        parquet_buffer = io.BytesIO()
                        st.session_state.df_principal.to_parquet(parquet_buffer, index=False)
                        # Guardamos como .parquet
                        zf.writestr("data.parquet", parquet_buffer.getvalue())
                    
                    # 2. Guardar Items Reporte (Pickle es necesario para objetos complejos como BytesIO de imágenes, 
                    # pero intentaremos serializar metadatos y guardar imágenes aparte para ser más limpios.
                    # Por compatibilidad rápida, usaremos pickle para la lista de dicts, pero es un riesgo menor si es local.
                    # MEJORA: Serializar reporte de forma segura.
                    # Por ahora mantenemos pickle para la lista de reportes por la complejidad de los objetos BytesIO dentro.)
                    if st.session_state.get('reporte_items'):
                        rep_buffer = io.BytesIO()
                        pickle.dump(st.session_state['reporte_items'], rep_buffer)
                        zf.writestr("report.pkl", rep_buffer.getvalue())
                    
                    # 3. Metadatos
                    meta = {
                        "fecha": str(datetime.datetime.now()),
                        "version": "2.0_secure",
                        "usuario": st.session_state.get('usuario_actual', 'Anonimo')
                    }
                    zf.writestr("metadata.json", json.dumps(meta, indent=4))
                
                zip_buffer.seek(0)
                st.download_button("⬇️ Descargar Proyecto (.zip)", zip_buffer, "proyecto_biostat.zip", "application/zip")
                st.success("Backup generado exitosamente.")
                
            except Exception as e:
                st.error(f"Error al guardar: {e}")

    # --- CARGAR ---
    with tab2:
        uploaded_zip = st.file_uploader("Cargar Backup (.zip)", type=["zip"])
        
        if uploaded_zip:
            if st.button("Restaurar Proyecto"):
                try:
                    with zipfile.ZipFile(uploaded_zip, "r") as zf:
                        # Listar archivos ignorando basura de Mac
                        files = [f for f in zf.namelist() if not f.startswith("__MACOSX") and not ".DS_Store" in f]
                        
                        # 1. Cargar Datos
                        # Intentar Parquet primero
                        if "data.parquet" in files:
                            with zf.open("data.parquet") as f:
                                st.session_state.df_principal = pd.read_parquet(f)
                                st.success("Datos (Parquet) restaurados.")
                        # Fallback a Pickle (Legacy)
                        elif "data.pkl" in files:
                            with zf.open("data.pkl") as f:
                                st.session_state.df_principal = pickle.load(f)
                                st.warning("Datos restaurados desde formato antiguo (Pickle). Se guardarán como Parquet la próxima vez.")
                        
                        # 2. Cargar Reporte
                        if "report.pkl" in files:
                            with zf.open("report.pkl") as f:
                                st.session_state['reporte_items'] = pickle.load(f)
                                st.success("Reporte restaurado.")
                        
                        st.rerun()
                        
                except Exception as e:
                    st.error(f"Error al restaurar: {e}")

    st.markdown('</div>', unsafe_allow_html=True)
